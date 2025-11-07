import sys
from PyQt5.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QMessageBox,
    QScrollArea, QGridLayout, QTextEdit, QGroupBox, QFileDialog, QSpinBox
)
from PyQt5.QtCore import Qt
from docx import Document
from docx.shared import Inches
from datetime import date
import os
import re

# --- Konfigurasi File ---
TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates")
TEMPLATE_FILENAME = "New_Template2.docx" 
TEMPLATE_PATH = os.path.join(TEMPLATES_DIR, TEMPLATE_FILENAME)

# --- Definisi Placeholders & Input Fields ---
FIELD_DEFINITIONS = {
    # Header fields
    "NO_TEST": {"placeholder": "[NO_TEST]", "label": "Nº de Test Plan", "type": "text"},
    "REV": {"placeholder": "[REV]", "label": "Revisión", "type": "text"},
    "DATE": {"placeholder": "[DATE]", "label": "Fecha de emisión (DD/MM/YYYY)", "type": "text"},

    # 0. INFORMACIÓN DEL SOLICITANTE DEL ENSAYO
    "TEXT1": {"placeholder": "[TEXT1]", "label": "Solicitante", "type": "text"},
    "TEXT4": {"placeholder": "[TEXT4]", "label": "Operario del ensayo", "type": "text"},
    "TEXT2": {"placeholder": "[TEXT2]", "label": "Departamento", "type": "text"},
    "TEXT5": {"placeholder": "[TEXT5]", "label": "Responsable del ensayo", "type": "text"},
    "TEXT3": {"placeholder": "[TEXT3]", "label": "Fecha de solicitud (DD/MM/YYYY)", "type": "text"},

    # 1. INFORMACIÓN GENERAL DEL PRODUCTO
    "TEXT6": {"placeholder": "[TEXT6]", "label": "Referencia del modelo ensayado", "type": "text"},
    "TEXT7": {"placeholder": "[TEXT7]", "label": "Aplicación", "type": "text"},
    "TEXT8": {"placeholder": "[TEXT8]", "label": "Fuente de luz", "type": "text"},

    # 1.1. CONDICIONES DEL ENSAYO
    "TEXT9": {"placeholder": "[TEXT9]", "label": "Ensayo térmico realizado en", "type": "text"},
    "TEXT10": {"placeholder": "[TEXT10]", "label": "Temperatura de color ensayada (CCT)", "type": "text"},
    "TEXT11": {"placeholder": "[TEXT11]", "label": "Luminaria alimentada a", "type": "text"},

    # 2. EQUIPOS Y MÉTODOS UTILIZADOS (up to 12)
    "EQUIPO1": {"placeholder": "[EQUIPO1]", "label": "Equipo 1", "type": "text"},
    "MARCA1": {"placeholder": "[MARCA1]", "label": "Marca/Modelo 1", "type": "text"},
    "TIPO1": {"placeholder": "[TIPO1]", "label": "Tipo/Aplicación 1", "type": "text"},
    "FECHA1": {"placeholder": "[FECHA1]", "label": "Fecha de calibración 1", "type": "text"},
    "OBSER1": {"placeholder": "[OBSER1]", "label": "Observaciones 1", "type": "text"},
    "EQUIPO2": {"placeholder": "[EQUIPO2]", "label": "Equipo 2", "type": "text"},
    "MARCA2": {"placeholder": "[MARCA2]", "label": "Marca/Modelo 2", "type": "text"},
    "TIPO2": {"placeholder": "[TIPO2]", "label": "Tipo/Aplicación 2", "type": "text"},
    "FECHA2": {"placeholder": "[FECHA2]", "label": "Fecha de calibración 2", "type": "text"},
    "OBSER2": {"placeholder": "[OBSER2]", "label": "Observaciones 2", "type": "text"},
    "EQUIPO3": {"placeholder": "[EQUIPO3]", "label": "Equipo 3", "type": "text"},
    "MARCA3": {"placeholder": "[MARCA3]", "label": "Marca/Modelo 3", "type": "text"},
    "TIPO3": {"placeholder": "[TIPO3]", "label": "Tipo/Aplicación 3", "type": "text"},
    "FECHA3": {"placeholder": "[FECHA3]", "label": "Fecha de calibración 3", "type": "text"},
    "OBSER3": {"placeholder": "[OBSER3]", "label": "Observaciones 3", "type": "text"},
    "EQUIPO4": {"placeholder": "[EQUIPO4]", "label": "Equipo 4", "type": "text"},
    "MARCA4": {"placeholder": "[MARCA4]", "label": "Marca/Modelo 4", "type": "text"},
    "TIPO4": {"placeholder": "[TIPO4]", "label": "Tipo/Aplicación 4", "type": "text"},
    "FECHA4": {"placeholder": "[FECHA4]", "label": "Fecha de calibración 4", "type": "text"},
    "OBSER4": {"placeholder": "[OBSER4]", "label": "Observaciones 4", "type": "text"},
    "EQUIPO5": {"placeholder": "[EQUIPO5]", "label": "Equipo 5", "type": "text"},
    "MARCA5": {"placeholder": "[MARCA5]", "label": "Marca/Modelo 5", "type": "text"},
    "TIPO5": {"placeholder": "[TIPO5]", "label": "Tipo/Aplicación 5", "type": "text"},
    "FECHA5": {"placeholder": "[FECHA5]", "label": "Fecha de calibración 5", "type": "text"},
    "OBSER5": {"placeholder": "[OBSER5]", "label": "Observaciones 5", "type": "text"},
    "EQUIPO6": {"placeholder": "[EQUIPO6]", "label": "Equipo 6", "type": "text"},
    "MARCA6": {"placeholder": "[MARCA6]", "label": "Marca/Modelo 6", "type": "text"},
    "TIPO6": {"placeholder": "[TIPO6]", "label": "Tipo/Aplicación 6", "type": "text"},
    "FECHA6": {"placeholder": "[FECHA6]", "label": "Fecha de calibración 6", "type": "text"},
    "OBSER6": {"placeholder": "[OBSER6]", "label": "Observaciones 6", "type": "text"},
    "EQUIPO7": {"placeholder": "[EQUIPO7]", "label": "Equipo 7", "type": "text"},
    "MARCA7": {"placeholder": "[MARCA7]", "label": "Marca/Modelo 7", "type": "text"},
    "TIPO7": {"placeholder": "[TIPO7]", "label": "Tipo/Aplicación 7", "type": "text"},
    "FECHA7": {"placeholder": "[FECHA7]", "label": "Fecha de calibración 7", "type": "text"},
    "OBSER7": {"placeholder": "[OBSER7]", "label": "Observaciones 7", "type": "text"},
    "EQUIPO8": {"placeholder": "[EQUIPO8]", "label": "Equipo 8", "type": "text"},
    "MARCA8": {"placeholder": "[MARCA8]", "label": "Marca/Modelo 8", "type": "text"},
    "TIPO8": {"placeholder": "[TIPO8]", "label": "Tipo/Aplicación 8", "type": "text"},
    "FECHA8": {"placeholder": "[FECHA8]", "label": "Fecha de calibración 8", "type": "text"},
    "OBSER8": {"placeholder": "[OBSER8]", "label": "Observaciones 8", "type": "text"},
    "EQUIPO9": {"placeholder": "[EQUIPO9]", "label": "Equipo 9", "type": "text"},
    "MARCA9": {"placeholder": "[MARCA9]", "label": "Marca/Modelo 9", "type": "text"},
    "TIPO9": {"placeholder": "[TIPO9]", "label": "Tipo/Aplicación 9", "type": "text"},
    "FECHA9": {"placeholder": "[FECHA9]", "label": "Fecha de calibración 9", "type": "text"},
    "OBSER9": {"placeholder": "[OBSER9]", "label": "Observaciones 9", "type": "text"},
    "EQUIPO10": {"placeholder": "[EQUIPO10]", "label": "Equipo 10", "type": "text"},
    "MARCA10": {"placeholder": "[MARCA10]", "label": "Marca/Modelo 10", "type": "text"},
    "TIPO10": {"placeholder": "[TIPO10]", "label": "Tipo/Aplicación 10", "type": "text"},
    "FECHA10": {"placeholder": "[FECHA10]", "label": "Fecha de calibración 10", "type": "text"},
    "OBSER10": {"placeholder": "[OBSER10]", "label": "Observaciones 10", "type": "text"},
    "EQUIPO11": {"placeholder": "[EQUIPO11]", "label": "Equipo 11", "type": "text"},
    "MARCA11": {"placeholder": "[MARCA11]", "label": "Marca/Modelo 11", "type": "text"},
    "TIPO11": {"placeholder": "[TIPO11]", "label": "Tipo/Aplicación 11", "type": "text"},
    "FECHA11": {"placeholder": "[FECHA11]", "label": "Fecha de calibración 11", "type": "text"},
    "OBSER11": {"placeholder": "[OBSER11]", "label": "Observaciones 11", "type": "text"},
    "EQUIPO12": {"placeholder": "[EQUIPO12]", "label": "Equipo 12", "type": "text"},
    "MARCA12": {"placeholder": "[MARCA12]", "label": "Marca/Modelo 12", "type": "text"},
    "TIPO12": {"placeholder": "[TIPO12]", "label": "Tipo/Aplicación 12", "type": "text"},
    "FECHA12": {"placeholder": "[FECHA12]", "label": "Fecha de calibración 12", "type": "text"},
    "OBSER12": {"placeholder": "[OBSER12]", "label": "Observaciones 12", "type": "text"},
    "TEXT12": {"placeholder": "[TEXT12]", "label": "Método de ensayo", "type": "text"},

    # 3. TEMPERATURAS REGISTRADAS (up to 10)
    "PUNTO1": {"placeholder": "[PUNTO1]", "label": "Punto de Medición 1", "type": "text"},
    "UNIDAD1": {"placeholder": "[UNIDAD1]", "label": "Unidad 1", "type": "text"},
    "LIMITE1": {"placeholder": "[LIMITE1]", "label": "Límite Máximo 1", "type": "text"},
    "TEMP1": {"placeholder": "[TEMP1]", "label": "Temperatura Medida 1", "type": "text"},
    "PUNTO2": {"placeholder": "[PUNTO2]", "label": "Punto de Medición 2", "type": "text"},
    "UNIDAD2": {"placeholder": "[UNIDAD2]", "label": "Unidad 2", "type": "text"},
    "LIMITE2": {"placeholder": "[LIMITE2]", "label": "Límite Máximo 2", "type": "text"},
    "TEMP2": {"placeholder": "[TEMP2]", "label": "Temperatura Medida 2", "type": "text"},
    "PUNTO3": {"placeholder": "[PUNTO3]", "label": "Punto de Medición 3", "type": "text"},
    "UNIDAD3": {"placeholder": "[UNIDAD3]", "label": "Unidad 3", "type": "text"},
    "LIMITE3": {"placeholder": "[LIMITE3]", "label": "Límite Máximo 3", "type": "text"},
    "TEMP3": {"placeholder": "[TEMP3]", "label": "Temperatura Medida 3", "type": "text"},
    "PUNTO4": {"placeholder": "[PUNTO4]", "label": "Punto de Medición 4", "type": "text"},
    "UNIDAD4": {"placeholder": "[UNIDAD4]", "label": "Unidad 4", "type": "text"},
    "LIMITE4": {"placeholder": "[LIMITE4]", "label": "Límite Máximo 4", "type": "text"},
    "TEMP4": {"placeholder": "[TEMP4]", "label": "Temperatura Medida 4", "type": "text"},
    "PUNTO5": {"placeholder": "[PUNTO5]", "label": "Punto de Medición 5", "type": "text"},
    "UNIDAD5": {"placeholder": "[UNIDAD5]", "label": "Unidad 5", "type": "text"},
    "LIMITE5": {"placeholder": "[LIMITE5]", "label": "Límite Máximo 5", "type": "text"},
    "TEMP5": {"placeholder": "[TEMP5]", "label": "Temperatura Medida 5", "type": "text"},
    "PUNTO6": {"placeholder": "[PUNTO6]", "label": "Punto de Medición 6", "type": "text"},
    "UNIDAD6": {"placeholder": "[UNIDAD6]", "label": "Unidad 6", "type": "text"},
    "LIMITE6": {"placeholder": "[LIMITE6]", "label": "Límite Máximo 6", "type": "text"},
    "TEMP6": {"placeholder": "[TEMP6]", "label": "Temperatura Medida 6", "type": "text"},
    "PUNTO7": {"placeholder": "[PUNTO7]", "label": "Punto de Medición 7", "type": "text"},
    "UNIDAD7": {"placeholder": "[UNIDAD7]", "label": "Unidad 7", "type": "text"},
    "LIMITE7": {"placeholder": "[LIMITE7]", "label": "Límite Máximo 7", "type": "text"},
    "TEMP7": {"placeholder": "[TEMP7]", "label": "Temperatura Medida 7", "type": "text"},
    "PUNTO8": {"placeholder": "[PUNTO8]", "label": "Punto de Medición 8", "type": "text"},
    "UNIDAD8": {"placeholder": "[UNIDAD8]", "label": "Unidad 8", "type": "text"},
    "LIMITE8": {"placeholder": "[LIMITE8]", "label": "Límite Máximo 8", "type": "text"},
    "TEMP8": {"placeholder": "[TEMP8]", "label": "Temperatura Medida 8", "type": "text"},
    "PUNTO9": {"placeholder": "[PUNTO9]", "label": "Punto de Medición 9", "type": "text"},
    "UNIDAD9": {"placeholder": "[UNIDAD9]", "label": "Unidad 9", "type": "text"},
    "LIMITE9": {"placeholder": "[LIMITE9]", "label": "Límite Máximo 9", "type": "text"},
    "TEMP9": {"placeholder": "[TEMP9]", "label": "Temperatura Medida 9", "type": "text"},
    "PUNTO10": {"placeholder": "[PUNTO10]", "label": "Punto de Medición 10", "type": "text"},
    "UNIDAD10": {"placeholder": "[UNIDAD10]", "label": "Unidad 10", "type": "text"},
    "LIMITE10": {"placeholder": "[LIMITE10]", "label": "Límite Máximo 10", "type": "text"},
    "TEMP10": {"placeholder": "[TEMP10]", "label": "Temperatura Medida 10", "type": "text"},
    "TEXT13": {"placeholder": "[TEXT13]", "label": "NOTA", "type": "text"},
    "TEXT14": {"placeholder": "[TEXT14]", "label": "Description", "type": "text"},
    "TEXT15": {"placeholder": "[TEXT15]", "label": "Conclusions", "type": "text"},

    # 3.1. GRÁFICA GENERADA
    "IMAGE1": {"placeholder": "[IMAGE1]", "label": "Imagen 1", "type": "file"},
    "TITLE1": {"placeholder": "[TITLE1]", "label": "Título 1", "type": "text"},
    "DESC1": {"placeholder": "[DESC1]", "label": "Descripción 1", "type": "text"},
    "IMAGE2": {"placeholder": "[IMAGE2]", "label": "Imagen 2", "type": "file"},
    "DESC2": {"placeholder": "[DESC2]", "label": "Descripción 2", "type": "text"},

    # 4. ESTABILIZACIÓN TÉRMICA (up to 10)
    "MEDICI1": {"placeholder": "[MEDICI1]", "label": "Punto de Medición 1", "type": "text"},
    "UNI1": {"placeholder": "[UNI1]", "label": "Unidad 1", "type": "text"},
    "VALMIN1": {"placeholder": "[VALMIN1]", "label": "Valor Mínimo 1", "type": "text"},
    "VALMAX1": {"placeholder": "[VALMAX1]", "label": "Valor Máximo 1", "type": "text"},
    "DESVI1": {"placeholder": "[DESVI1]", "label": "Desviación 1", "type": "text"},
    "MEDICI2": {"placeholder": "[MEDICI2]", "label": "Punto de Medición 2", "type": "text"},
    "UNI2": {"placeholder": "[UNI2]", "label": "Unidad 2", "type": "text"},
    "VALMIN2": {"placeholder": "[VALMIN2]", "label": "Valor Mínimo 2", "type": "text"},
    "VALMAX2": {"placeholder": "[VALMAX2]", "label": "Valor Máximo 2", "type": "text"},
    "DESVI2": {"placeholder": "[DESVI2]", "label": "Desviación 2", "type": "text"},
    "MEDICI3": {"placeholder": "[MEDICI3]", "label": "Punto de Medición 3", "type": "text"},
    "UNI3": {"placeholder": "[UNI3]", "label": "Unidad 3", "type": "text"},
    "VALMIN3": {"placeholder": "[VALMIN3]", "label": "Valor Mínimo 3", "type": "text"},
    "VALMAX3": {"placeholder": "[VALMAX3]", "label": "Valor Máximo 3", "type": "text"},
    "DESVI3": {"placeholder": "[DESVI3]", "label": "Desviación 3", "type": "text"},
    "MEDICI4": {"placeholder": "[MEDICI4]", "label": "Punto de Medición 4", "type": "text"},
    "UNI4": {"placeholder": "[UNI4]", "label": "Unidad 4", "type": "text"},
    "VALMIN4": {"placeholder": "[VALMIN4]", "label": "Valor Mínimo 4", "type": "text"},
    "VALMAX4": {"placeholder": "[VALMAX4]", "label": "Valor Máximo 4", "type": "text"},
    "DESVI4": {"placeholder": "[DESVI4]", "label": "Desviación 4", "type": "text"},
    "MEDICI5": {"placeholder": "[MEDICI5]", "label": "Punto de Medición 5", "type": "text"},
    "UNI5": {"placeholder": "[UNI5]", "label": "Unidad 5", "type": "text"},
    "VALMIN5": {"placeholder": "[VALMIN5]", "label": "Valor Mínimo 5", "type": "text"},
    "VALMAX5": {"placeholder": "[VALMAX5]", "label": "Valor Máximo 5", "type": "text"},
    "DESVI5": {"placeholder": "[DESVI5]", "label": "Desviación 5", "type": "text"},
    "MEDICI6": {"placeholder": "[MEDICI6]", "label": "Punto de Medición 6", "type": "text"},
    "UNI6": {"placeholder": "[UNI6]", "label": "Unidad 6", "type": "text"},
    "VALMIN6": {"placeholder": "[VALMIN6]", "label": "Valor Mínimo 6", "type": "text"},
    "VALMAX6": {"placeholder": "[VALMAX6]", "label": "Valor Máximo 6", "type": "text"},
    "DESVI6": {"placeholder": "[DESVI6]", "label": "Desviación 6", "type": "text"},
    "MEDICI7": {"placeholder": "[MEDICI7]", "label": "Punto de Medición 7", "type": "text"},
    "UNI7": {"placeholder": "[UNI7]", "label": "Unidad 7", "type": "text"},
    "VALMIN7": {"placeholder": "[VALMIN7]", "label": "Valor Mínimo 7", "type": "text"},
    "VALMAX7": {"placeholder": "[VALMAX7]", "label": "Valor Máximo 7", "type": "text"},
    "DESVI7": {"placeholder": "[DESVI7]", "label": "Desviación 7", "type": "text"},
    "MEDICI8": {"placeholder": "[MEDICI8]", "label": "Punto de Medición 8", "type": "text"},
    "UNI8": {"placeholder": "[UNI8]", "label": "Unidad 8", "type": "text"},
    "VALMIN8": {"placeholder": "[VALMIN8]", "label": "Valor Mínimo 8", "type": "text"},
    "VALMAX8": {"placeholder": "[VALMAX8]", "label": "Valor Máximo 8", "type": "text"},
    "DESVI8": {"placeholder": "[DESVI8]", "label": "Desviación 8", "type": "text"},
    "MEDICI9": {"placeholder": "[MEDICI9]", "label": "Punto de Medición 9", "type": "text"},
    "UNI9": {"placeholder": "[UNI9]", "label": "Unidad 9", "type": "text"},
    "VALMIN9": {"placeholder": "[VALMIN9]", "label": "Valor Mínimo 9", "type": "text"},
    "VALMAX9": {"placeholder": "[VALMAX9]", "label": "Valor Máximo 9", "type": "text"},
    "DESVI9": {"placeholder": "[DESVI9]", "label": "Desviación 9", "type": "text"},
    "MEDICI10": {"placeholder": "[MEDICI10]", "label": "Punto de Medición 10", "type": "text"},
    "UNI10": {"placeholder": "[UNI10]", "label": "Unidad 10", "type": "text"},
    "VALMIN10": {"placeholder": "[VALMIN10]", "label": "Valor Mínimo 10", "type": "text"},
    "VALMAX10": {"placeholder": "[VALMAX10]", "label": "Valor Máximo 10", "type": "text"},
    "DESVI10": {"placeholder": "[DESVI10]", "label": "Desviación 10", "type": "text"},

    # 5. RESULTADOS (up to 10)
    "PUNTODE1": {"placeholder": "[PUNTODE1]", "label": "Punto de Medición 1", "type": "text"},
    "UNIC1": {"placeholder": "[UNIC1]", "label": "Unidad 1", "type": "text"},
    "TEMPE1": {"placeholder": "[TEMPE1]", "label": "Temperatura final 1", "type": "text"},
    "RESULT1": {"placeholder": "[RESULT1]", "label": "Resultado 1", "type": "text"},
    "PUNTODE2": {"placeholder": "[PUNTODE2]", "label": "Punto de Medición 2", "type": "text"},
    "UNIC2": {"placeholder": "[UNIC2]", "label": "Unidad 2", "type": "text"},
    "TEMPE2": {"placeholder": "[TEMPE2]", "label": "Temperatura final 2", "type": "text"},
    "RESULT2": {"placeholder": "[RESULT2]", "label": "Resultado 2", "type": "text"},
    "PUNTODE3": {"placeholder": "[PUNTODE3]", "label": "Punto de Medición 3", "type": "text"},
    "UNIC3": {"placeholder": "[UNIC3]", "label": "Unidad 3", "type": "text"},
    "TEMPE3": {"placeholder": "[TEMPE3]", "label": "Temperatura final 3", "type": "text"},
    "RESULT3": {"placeholder": "[RESULT3]", "label": "Resultado 3", "type": "text"},
    "PUNTODE4": {"placeholder": "[PUNTODE4]", "label": "Punto de Medición 4", "type": "text"},
    "UNIC4": {"placeholder": "[UNIC4]", "label": "Unidad 4", "type": "text"},
    "TEMPE4": {"placeholder": "[TEMPE4]", "label": "Temperatura final 4", "type": "text"},
    "RESULT4": {"placeholder": "[RESULT4]", "label": "Resultado 4", "type": "text"},
    "PUNTODE5": {"placeholder": "[PUNTODE5]", "label": "Punto de Medición 5", "type": "text"},
    "UNIC5": {"placeholder": "[UNIC5]", "label": "Unidad 5", "type": "text"},
    "TEMPE5": {"placeholder": "[TEMPE5]", "label": "Temperatura final 5", "type": "text"},
    "RESULT5": {"placeholder": "[RESULT5]", "label": "Resultado 5", "type": "text"},
    "PUNTODE6": {"placeholder": "[PUNTODE6]", "label": "Punto de Medición 6", "type": "text"},
    "UNIC6": {"placeholder": "[UNIC6]", "label": "Unidad 6", "type": "text"},
    "TEMPE6": {"placeholder": "[TEMPE6]", "label": "Temperatura final 6", "type": "text"},
    "RESULT6": {"placeholder": "[RESULT6]", "label": "Resultado 6", "type": "text"},
    "PUNTODE7": {"placeholder": "[PUNTODE7]", "label": "Punto de Medición 7", "type": "text"},
    "UNIC7": {"placeholder": "[UNIC7]", "label": "Unidad 7", "type": "text"},
    "TEMPE7": {"placeholder": "[TEMPE7]", "label": "Temperatura final 7", "type": "text"},
    "RESULT7": {"placeholder": "[RESULT7]", "label": "Resultado 7", "type": "text"},
    "PUNTODE8": {"placeholder": "[PUNTODE8]", "label": "Punto de Medición 8", "type": "text"},
    "UNIC8": {"placeholder": "[UNIC8]", "label": "Unidad 8", "type": "text"},
    "TEMPE8": {"placeholder": "[TEMPE8]", "label": "Temperatura final 8", "type": "text"},
    "RESULT8": {"placeholder": "[RESULT8]", "label": "Resultado 8", "type": "text"},
    "PUNTODE9": {"placeholder": "[PUNTODE9]", "label": "Punto de Medición 9", "type": "text"},
    "UNIC9": {"placeholder": "[UNIC9]", "label": "Unidad 9", "type": "text"},
    "TEMPE9": {"placeholder": "[TEMPE9]", "label": "Temperatura final 9", "type": "text"},
    "RESULT9": {"placeholder": "[RESULT9]", "label": "Resultado 9", "type": "text"},
    "PUNTODE10": {"placeholder": "[PUNTODE10]", "label": "Punto de Medición 10", "type": "text"},
    "UNIC10": {"placeholder": "[UNIC10]", "label": "Unidad 10", "type": "text"},
    "TEMPE10": {"placeholder": "[TEMPE10]", "label": "Temperatura final 10", "type": "text"},
    "RESULT10": {"placeholder": "[RESULT10]", "label": "Resultado 10", "type": "text"},

    # 7. FOTOGRAFIAS (up to 10)
    "IMAGE3": {"placeholder": "[IMAGE3]", "label": "Imagen 3", "type": "file"},
    "TITLE3": {"placeholder": "[TITLE3]", "label": "Título 3", "type": "text"},
    "IMAGE4": {"placeholder": "[IMAGE4]", "label": "Imagen 4", "type": "file"},
    "TITLE4": {"placeholder": "[TITLE4]", "label": "Título 4", "type": "text"},
    "IMAGE5": {"placeholder": "[IMAGE5]", "label": "Imagen 5", "type": "file"},
    "TITLE5": {"placeholder": "[TITLE5]", "label": "Título 5", "type": "text"},
    "IMAGE6": {"placeholder": "[IMAGE6]", "label": "Imagen 6", "type": "file"},
    "TITLE6": {"placeholder": "[TITLE6]", "label": "Título 6", "type": "text"},
    "IMAGE7": {"placeholder": "[IMAGE7]", "label": "Imagen 7", "type": "file"},
    "TITLE7": {"placeholder": "[TITLE7]", "label": "Título 7", "type": "text"},
    "IMAGE8": {"placeholder": "[IMAGE8]", "label": "Imagen 8", "type": "file"},
    "TITLE8": {"placeholder": "[TITLE8]", "label": "Título 8", "type": "text"},
    "TITLE9": {"placeholder": "[TITLE9]", "label": "Título 9", "type": "text"},
    "IMAGE9": {"placeholder": "[IMAGE9]", "label": "Imagen 9", "type": "file"},
    "TITLE10": {"placeholder": "[TITLE10]", "label": "Título 10", "type": "text"},
    "IMAGE10": {"placeholder": "[IMAGE10]", "label": "Imagen 10", "type": "file"},
    "TITLE11": {"placeholder": "[TITLE11]", "label": "Título 11", "type": "text"},
    "IMAGE11": {"placeholder": "[IMAGE11]", "label": "Imagen 11", "type": "file"},
    "TITLE12": {"placeholder": "[TITLE12]", "label": "Título 12", "type": "text"},
    "IMAGE12": {"placeholder": "[IMAGE12]", "label": "Imagen 12", "type": "file"},
}


class DocumentGeneratorApp(QWidget):
    """Aplikasi untuk menginput data dan menghasilkan dokumen Word dari template."""
    def __init__(self):
        super().__init__()
        self.input_widgets = {}
        self.equipment_groups = []
        self.spin_boxes = {}
        self.setWindowTitle("Generador de Anexo II al Informe")
        self.setStyleSheet("font-size: 14px; font-family: Arial;")
        self.init_ui()

    def init_ui(self):
        """Membangun antarmuka pengguna."""
        main_layout = QVBoxLayout(self)

        # --- Judul ---
        title = QLabel("Ingresar Datos Para el Anexo II")
        title.setAlignment(Qt.AlignCenter)
        title.setStyleSheet("font-size: 20px; font-weight: bold; margin-bottom: 10px; color: #2C3E50;")
        main_layout.addWidget(title)

        # --- Spin Boxes for Row Selection ---
        spin_layout = QHBoxLayout()
        spin_layout.addWidget(QLabel("EQUIPOS Y MÉTODOS UTILIZADOS (max 12):"))
        self.spin_equipment = QSpinBox()
        self.spin_equipment.setRange(1, 12)
        self.spin_equipment.setValue(12)
        self.spin_equipment.valueChanged.connect(self.rebuild_form)
        spin_layout.addWidget(self.spin_equipment)

        spin_layout.addWidget(QLabel("TEMPERATURAS REGISTRADAS (max 10):"))
        self.spin_temperatures = QSpinBox()
        self.spin_temperatures.setRange(1, 10)
        self.spin_temperatures.setValue(10)
        self.spin_temperatures.valueChanged.connect(self.rebuild_form)
        spin_layout.addWidget(self.spin_temperatures)

        spin_layout.addWidget(QLabel("ESTABILIZACIÓN TÉRMICA (max 10):"))
        self.spin_stabilization = QSpinBox()
        self.spin_stabilization.setRange(1, 10)
        self.spin_stabilization.setValue(10)
        self.spin_stabilization.valueChanged.connect(self.rebuild_form)
        spin_layout.addWidget(self.spin_stabilization)

        spin_layout.addWidget(QLabel("RESULTADOS (max 10):"))
        self.spin_results = QSpinBox()
        self.spin_results.setRange(1, 10)
        self.spin_results.setValue(10)
        self.spin_results.valueChanged.connect(self.rebuild_form)
        spin_layout.addWidget(self.spin_results)

        main_layout.addLayout(spin_layout)

        # --- Scroll Area untuk banyak input ---
        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.content_widget = QWidget()
        self.form_layout = QVBoxLayout(self.content_widget)
        self.form_layout.setSpacing(15)
        self.scroll.setWidget(self.content_widget)
        main_layout.addWidget(self.scroll)

        # --- Tombol Generate ---
        self.generate_button = QPushButton("GENERAR DOCUMENTO DE WORD (.docx)")
        self.generate_button.setStyleSheet(
            "background-color: #3498DB; color: white; padding: 12px; border-radius: 8px; font-weight: bold;"
        )
        self.generate_button.clicked.connect(self.generate_document)
        main_layout.addWidget(self.generate_button)

        # --- Informasi Template ---
        info = QLabel(
            f"**Plantilla utilizada:** '{TEMPLATE_FILENAME}'\n"
            f"Asegúrate de que este archivo esté en la carpeta: '{TEMPLATES_DIR}'"
        )
        info.setStyleSheet("font-size: 10px; color: gray; margin-top: 5px;")
        main_layout.addWidget(info)

        self.setLayout(main_layout)
        self.resize(600, 700)
        self.rebuild_form()

    def rebuild_form(self):
        # Clear existing widgets from form_layout
        for i in reversed(range(self.form_layout.count())):
            item = self.form_layout.itemAt(i)
            if item.widget():
                item.widget().setParent(None)
            elif item.layout():
                # If it's a layout, remove it
                sub_layout = item.layout()
                while sub_layout.count():
                    sub_item = sub_layout.takeAt(0)
                    if sub_item.widget():
                        sub_item.widget().setParent(None)
                self.form_layout.removeItem(item)

        self.input_widgets = {}

        # Header
        title_label = QLabel("Encabezado - Información del documento")
        title_label.setStyleSheet("font-size: 18px; font-weight: bold; margin-bottom: 10px;")
        self.form_layout.addWidget(title_label)
        self.create_input_group(self.form_layout, "Encabezado - Información del documento", [
            "NO_TEST", "REV", "DATE"
        ])

        # 0. INFORMACIÓN DEL SOLICITANTE DEL ENSAYO
        title_label = QLabel("0. INFORMACIÓN DEL SOLICITANTE DEL ENSAYO")
        title_label.setStyleSheet("font-size: 18px; font-weight: bold; margin-bottom: 10px;")
        self.form_layout.addWidget(title_label)
        self.create_input_group(self.form_layout, "0. INFORMACIÓN DEL SOLICITANTE DEL ENSAYO", [
            "TEXT1", "TEXT4", "TEXT2", "TEXT5", "TEXT3"
        ])

        # 1. INFORMACIÓN GENERAL DEL PRODUCTO
        title_label = QLabel("1. INFORMACIÓN GENERAL DEL PRODUCTO")
        title_label.setStyleSheet("font-size: 18px; font-weight: bold; margin-bottom: 10px;")
        self.form_layout.addWidget(title_label)
        self.create_input_group(self.form_layout, "1. INFORMACIÓN GENERAL DEL PRODUCTO", [
            "TEXT6", "TEXT7", "TEXT8"
        ])

        # 1.1. CONDICIONES DEL ENSAYO
        self.create_input_group(self.form_layout, "1.1. CONDICIONES DEL ENSAYO", [
            "TEXT9", "TEXT10", "TEXT11"
        ])

        # 2. EQUIPOS Y MÉTODOS UTILIZADOS
        title_label = QLabel("2. EQUIPOS Y MÉTODOS UTILIZADOS")
        title_label.setStyleSheet("font-size: 18px; font-weight: bold; margin-bottom: 10px;")
        self.form_layout.addWidget(title_label)
        num_equip = self.spin_equipment.value()
        for i in range(1, num_equip + 1):
            self.create_input_group(self.form_layout, f"Row {i}", [
                f"EQUIPO{i}", f"MARCA{i}", f"TIPO{i}", f"FECHA{i}", f"OBSER{i}"
            ])

        # 2.1. MÉTODO DE ENSAYO
        self.create_input_group(self.form_layout, "MÉTODO DE ENSAYO", [
            "TEXT12"
        ])

        # 3. TEMPERATURAS REGISTRADAS
        title_label = QLabel("3. TEMPERATURAS REGISTRADAS")
        title_label.setStyleSheet("font-size: 18px; font-weight: bold; margin-bottom: 10px;")
        self.form_layout.addWidget(title_label)
        num_temp = self.spin_temperatures.value()
        for i in range(1, num_temp + 1):
            self.create_input_group(self.form_layout, f"Row {i} ", [
                f"PUNTO{i}", f"UNIDAD{i}", f"LIMITE{i}", f"TEMP{i}"
            ])
        self.create_input_group(self.form_layout, "NOTA", [
            "TEXT13"
        ])

        # 3.1. GRÁFICA GENERADA
        self.create_input_group(self.form_layout, "3.1. GRÁFICA GENERADA", [
            "IMAGE1", "TITLE1", "DESC1", "IMAGE2", "DESC2"
        ])

        # 4. ESTABILIZACIÓN TÉRMICA
        title_label = QLabel("4. ESTABILIZACIÓN TÉRMICA")
        title_label.setStyleSheet("font-size: 18px; font-weight: bold; margin-bottom: 10px;")
        self.form_layout.addWidget(title_label)
        self.create_input_group(self.form_layout, "Description", [
            "TEXT14"
        ])
        num_stab = self.spin_stabilization.value()
        for i in range(1, num_stab + 1):
            self.create_input_group(self.form_layout, f"Row {i}", [
                f"MEDICI{i}", f"UNI{i}", f"VALMIN{i}", f"VALMAX{i}", f"DESVI{i}"
            ])

        # 5. RESULTADOS
        title_label = QLabel("5. RESULTADOS")
        title_label.setStyleSheet("font-size: 18px; font-weight: bold; margin-bottom: 10px;")
        self.form_layout.addWidget(title_label)
        num_res = self.spin_results.value()
        for i in range(1, num_res + 1):
            self.create_input_group(self.form_layout, f"Row {i}", [
                f"PUNTODE{i}", f"UNIC{i}", f"TEMPE{i}", f"RESULT{i}"
            ])

        # 6. CONCLUSIONES DEL LABORATORIO
        title_label = QLabel("6. CONCLUSIONES DEL LABORATORIO")
        title_label.setStyleSheet("font-size: 18px; font-weight: bold; margin-bottom: 10px;")
        self.form_layout.addWidget(title_label)
        self.create_input_group(self.form_layout, "CONCLUSIONES", [
            "TEXT15"
        ])

        # 7. FOTOGRAFIAS
        title_label = QLabel("7. FOTOGRAFIAS")
        title_label.setStyleSheet("font-size: 18px; font-weight: bold; margin-bottom: 10px;")
        self.form_layout.addWidget(title_label)
        for i in range(3, 13):  # IMAGE3 to IMAGE12, TITLE3 to TITLE12
            self.create_input_group(self.form_layout, f"Fotografía {i-2}", [
                f"IMAGE{i}"
            ])
            self.create_input_group(self.form_layout, f"Titulo {i-2}", [
                f"TITLE{i}"
            ])

    def create_input_group(self, parent_layout, title, keys):
        """Membuat group box untuk input yang terorganisir."""
        group_box = QGroupBox(title)
        group_box.setStyleSheet("font-weight: bold; margin-top: 10px;")
        grid_layout = QGridLayout()
        grid_layout.setSpacing(10)
        
        row = 0
        col = 0
        
        for key in keys:
            definition = FIELD_DEFINITIONS[key]
            
            label = QLabel(f"{definition['label']}:")
            
            if definition['type'] == "text":
                if key in ["TEXT1", "TEXT4", "TEXT2", "TEXT5", "TEXT3", "TEXT6", "TEXT7", "TEXT8", "TEXT9", "TEXT10", "TEXT11", "TEXT12", "TEXT13", "TEXT14", "TEXT15"]:
                    input_field = QTextEdit()
                    input_field.setMinimumHeight(60)
                    grid_layout.addWidget(label, row, 0, 1, 2)
                    grid_layout.addWidget(input_field, row + 1, 0, 1, 2)
                    row += 2
                    col = 0
                else:
                    input_field = QLineEdit()
                    input_field.setMinimumHeight(30)
                    grid_layout.addWidget(label, row, col)
                    grid_layout.addWidget(input_field, row + 1, col)
                    col = 1 - col
                    if col == 0:
                        row += 2
            elif definition['type'] == "file":
                input_field = QLineEdit()
                input_field.setMinimumHeight(30)
                browse_button = QPushButton("Browse")
                browse_button.clicked.connect(lambda _, field=input_field: self.browse_file(field))
                grid_layout.addWidget(label, row, 0)
                grid_layout.addWidget(input_field, row + 1, 0)
                grid_layout.addWidget(browse_button, row + 1, 1)
                row += 2
                col = 0
            
            self.input_widgets[key] = input_field

        group_box.setLayout(grid_layout)
        parent_layout.addWidget(group_box)
        
    def browse_file(self, field):
        """Browse file untuk input gambar."""
        file_path, _ = QFileDialog.getOpenFileName(self, "Select Image", "", "Image Files (*.png *.jpg *.jpeg *.gif *.bmp *.tiff *.tif *.webp *.jfif)")
        if file_path:
            field.setText(file_path)

    def replace_in_paragraph(self, paragraph, placeholder, value):
        """Mengganti placeholder di dalam paragraf."""
        if placeholder in paragraph.text and 'IMAGE' not in placeholder:
            paragraph.text = paragraph.text.replace(placeholder, value)
            for run in paragraph.runs:
                run.font.name = "Gordita Light"
                run.font.size = 9 * 12700

    def replace_in_tables(self, document, replacement_data):
        """Mengganti placeholder di dalam sel tabel."""
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in replacement_data.items():
                            self.replace_in_paragraph(paragraph, placeholder, value)

    def replace_in_headers(self, document, replacement_data):
        """Mengganti placeholder di dalam header."""
        for section in document.sections:
            header = section.header
            for paragraph in header.paragraphs:
                for placeholder, value in replacement_data.items():
                    self.replace_in_paragraph(paragraph, placeholder, value)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, value in replacement_data.items():
                                self.replace_in_paragraph(paragraph, placeholder, value)

    def replace_in_footers(self, document, replacement_data):
        """Mengganti placeholder di dalam footer."""
        for section in document.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                for placeholder, value in replacement_data.items():
                    self.replace_in_paragraph(paragraph, placeholder, value)
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, value in replacement_data.items():
                                self.replace_in_paragraph(paragraph, placeholder, value)

    def replace_images(self, document, replacement_data):
        """Mengganti placeholder gambar dengan gambar yang dipilih."""
        # Replace in main paragraphs
        for paragraph in document.paragraphs:
            placeholder = paragraph.text.strip()
            if placeholder in replacement_data:
                image_path = replacement_data[placeholder]
                if os.path.exists(image_path):
                    paragraph.clear()
                    # Check if it's a photography image (IMAGE3 to IMAGE8) for vertical orientation
                    if placeholder in ["[IMAGE3]", "[IMAGE4]", "[IMAGE5]", "[IMAGE6]", "[IMAGE7]", "[IMAGE8]"]:
                        # Vertical photo from phone: height 4 inches, width 3 inches
                        paragraph.add_run().add_picture(image_path, width=Inches(3), height=Inches(4))
                    else:
                        # Other images: default size
                        paragraph.add_run().add_picture(image_path)

        # Replace in tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        placeholder = paragraph.text.strip()
                        if placeholder in replacement_data:
                            image_path = replacement_data[placeholder]
                            if os.path.exists(image_path):
                                paragraph.clear()
                                # Check if it's a photography image (IMAGE3 to IMAGE8) for vertical orientation
                                if placeholder in ["[IMAGE3]", "[IMAGE4]", "[IMAGE5]", "[IMAGE6]", "[IMAGE7]", "[IMAGE8]"]:
                                    # Vertical photo from phone: height 4 inches, width 3 inches
                                    paragraph.add_run().add_picture(image_path, width=Inches(3), height=Inches(4))
                                else:
                                    # Other images: default size
                                    paragraph.add_run().add_picture(image_path)

        # Replace in headers
        for section in document.sections:
            header = section.header
            for paragraph in header.paragraphs:
                placeholder = paragraph.text.strip()
                if placeholder in replacement_data:
                    image_path = replacement_data[placeholder]
                    if os.path.exists(image_path):
                        paragraph.clear()
                        # Check if it's a photography image (IMAGE3 to IMAGE8) for vertical orientation
                        if placeholder in ["[IMAGE3]", "[IMAGE4]", "[IMAGE5]", "[IMAGE6]", "[IMAGE7]", "[IMAGE8]"]:
                            # Vertical photo from phone: height 4 inches, width 3 inches
                            paragraph.add_run().add_picture(image_path, width=Inches(3), height=Inches(4))
                        else:
                            # Other images: default size
                            paragraph.add_run().add_picture(image_path)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            placeholder = paragraph.text.strip()
                            if placeholder in replacement_data:
                                image_path = replacement_data[placeholder]
                                if os.path.exists(image_path):
                                    paragraph.clear()
                                    # Check if it's a photography image (IMAGE3 to IMAGE8) for vertical orientation
                                    if placeholder in ["[IMAGE3]", "[IMAGE4]", "[IMAGE5]", "[IMAGE6]", "[IMAGE7]", "[IMAGE8]"]:
                                        # Vertical photo from phone: height 4 inches, width 3 inches
                                        paragraph.add_run().add_picture(image_path, width=Inches(3), height=Inches(4))
                                    else:
                                        # Other images: default size
                                        paragraph.add_run().add_picture(image_path)

        # Replace in footers
        for section in document.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                placeholder = paragraph.text.strip()
                if placeholder in replacement_data:
                    image_path = replacement_data[placeholder]
                    if os.path.exists(image_path):
                        paragraph.clear()
                        # Check if it's a photography image (IMAGE3 to IMAGE8) for vertical orientation
                        if placeholder in ["[IMAGE3]", "[IMAGE4]", "[IMAGE5]", "[IMAGE6]", "[IMAGE7]", "[IMAGE8]"]:
                            # Vertical photo from phone: height 4 inches, width 3 inches
                            paragraph.add_run().add_picture(image_path, width=Inches(3), height=Inches(4))
                        else:
                            # Other images: default size
                            paragraph.add_run().add_picture(image_path)
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            placeholder = paragraph.text.strip()
                            if placeholder in replacement_data:
                                image_path = replacement_data[placeholder]
                                if os.path.exists(image_path):
                                    paragraph.clear()
                                    # Check if it's a photography image (IMAGE3 to IMAGE8) for vertical orientation
                                    if placeholder in ["[IMAGE3]", "[IMAGE4]", "[IMAGE5]", "[IMAGE6]", "[IMAGE7]", "[IMAGE8]"]:
                                        # Vertical photo from phone: height 4 inches, width 3 inches
                                        paragraph.add_run().add_picture(image_path, width=Inches(3), height=Inches(4))
                                    else:
                                        # Other images: default size
                                        paragraph.add_run().add_picture(image_path)

    def generate_document(self):
        """Logika utama untuk membaca input, memuat template, mengganti placeholder, dan menyimpan file."""

        replacement_data = {}
        all_required_filled = True

        # Collect data from used widgets
        for key in self.input_widgets:
            definition = FIELD_DEFINITIONS[key]
            input_widget = self.input_widgets[key]
            if isinstance(input_widget, QLineEdit):
                value = input_widget.text().strip()
            elif isinstance(input_widget, QTextEdit):
                value = input_widget.toPlainText().strip()
            else:
                continue

            if key in ["TEXT1", "TEXT4", "TEXT2", "TEXT5", "TEXT3", "TEXT6", "TEXT7", "TEXT8", "TEXT9", "TEXT10", "TEXT11", "TEXT12", "TEXT13", "TEXT14", "TEXT15"] and not value:
                all_required_filled = False
                QMessageBox.warning(self, "Input Kosong", f"Campo obligatorio ('{definition['label']}') no puede estar vacío.")
                return

            replacement_data[definition['placeholder']] = value

        # Replace unused placeholders with empty strings to remove them
        for key, definition in FIELD_DEFINITIONS.items():
            if key not in self.input_widgets:
                replacement_data[definition['placeholder']] = ""
        
        if not all_required_filled:
            return

        if not os.path.exists(TEMPLATE_PATH):
            QMessageBox.critical(self, "Error", 
                f"Plantilla no encontrada en: {TEMPLATE_PATH}. "
                f"Por favor coloque el archivo '{TEMPLATE_FILENAME}' que usted proporciona en la carpeta 'templates'."
            )
            return

        try:
            document = Document(TEMPLATE_PATH)
        except Exception as e:
            QMessageBox.critical(self, "Error al leer la plantilla", f"Error al cargar la plantilla: {e}")
            return

        for paragraph in document.paragraphs:
            for placeholder, value in replacement_data.items():
                self.replace_in_paragraph(paragraph, placeholder, value)

        self.replace_in_tables(document, replacement_data)
        self.replace_in_headers(document, replacement_data)
        self.replace_in_footers(document, replacement_data)
        self.replace_images(document, replacement_data)

        # Remove empty table rows
        for table in document.tables:
            rows_to_remove = []
            for i, row in enumerate(table.rows):
                is_empty = True
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            is_empty = False
                            break
                    if not is_empty:
                        break
                if is_empty:
                    rows_to_remove.append(i)
            for i in reversed(rows_to_remove):
                table._tbl.remove(table.rows[i]._tr)

        try:
            output_filename = f"Generated_Anexo_II_{date.today().strftime('%d_%m_%Y')}.docx"
            file_path, _ = QFileDialog.getSaveFileName(self, "Guardar documento como...", output_filename, "Word Documents (*.docx);;All Files (*)", options=QFileDialog.Options())
            if not file_path:
                QMessageBox.information(self, "Cancelado", "El almacenamiento fue cancelado por el usuario.")
                return
            if not file_path.lower().endswith('.docx'):
                file_path += '.docx'

            document.save(file_path)
            QMessageBox.information(
                self,
                "¡Listo!",
                f"El documento de Word se creó y se guardó con éxito como:\n{file_path}"
            )
        except Exception as e:
            QMessageBox.critical(self, "Error al guardar el archivo", f"Error al guardar el documento: {e}")


if __name__ == '__main__':
    if not os.path.exists(TEMPLATES_DIR):
        os.makedirs(TEMPLATES_DIR)
        print(f"La carpeta 'templates' acaba de ser creada. Por favor, coloque el archivo '{TEMPLATE_FILENAME}' Dentro de él, luego vuelve a ejecutar la aplicación.")
        sys.exit()

    app = QApplication(sys.argv)
    window = DocumentGeneratorApp()
    window.show()
    sys.exit(app.exec_())
