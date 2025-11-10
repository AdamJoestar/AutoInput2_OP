import os
from PyQt5.QtWidgets import QMessageBox, QFileDialog, QComboBox, QDateEdit
from docx import Document
from docx.shared import Inches
from datetime import date
from fields import FIELD_DEFINITIONS
from utils import validate_required_fields


class DocumentProcessor:
    def __init__(self, parent_app):
        self.parent_app = parent_app

    def generate_document(self):
        """Logika utama untuk membaca input, memuat template, mengganti placeholder, dan menyimpan file."""
        input_widgets = self.parent_app.ui_builder.input_widgets

        # Validate required fields
        is_valid, missing_field = validate_required_fields(input_widgets, FIELD_DEFINITIONS)
        if not is_valid:
            QMessageBox.warning(self.parent_app, "Input Kosong", f"Campo obligatorio ('{missing_field}') no puede estar vacío.")
            return

        replacement_data = self.collect_replacement_data(input_widgets)

        # Check template existence
        if not os.path.exists(self.parent_app.template_path):
            QMessageBox.critical(self.parent_app, "Error", 
                f"Plantilla no encontrada en: {self.parent_app.template_path}. "
                f"Por favor coloque el archivo '{self.parent_app.template_filename}' que usted proporciona en la carpeta 'templates'."
            )
            return

        try:
            document = Document(self.parent_app.template_path)
        except Exception as e:
            QMessageBox.critical(self.parent_app, "Error al leer la plantilla", f"Error al cargar la plantilla: {e}")
            return

        self.process_document(document, replacement_data)
        self.save_document(document)

    def collect_replacement_data(self, input_widgets):
        replacement_data = {}
        for key in input_widgets:
            definition = FIELD_DEFINITIONS[key]
            input_widget = input_widgets[key]
            if isinstance(input_widget, QComboBox):
                value = input_widget.currentText().strip()
            elif isinstance(input_widget, QDateEdit):
                value = input_widget.date().toString("dd/MM/yyyy")
            elif hasattr(input_widget, 'text'):
                value = input_widget.text().strip()
            elif hasattr(input_widget, 'toPlainText'):
                value = input_widget.toPlainText().strip()
            else:
                continue
            replacement_data[definition['placeholder']] = value

        # Replace unused placeholders with empty strings
        for key, definition in FIELD_DEFINITIONS.items():
            if key not in input_widgets:
                replacement_data[definition['placeholder']] = ""
        return replacement_data

    def process_document(self, document, replacement_data):
        for paragraph in document.paragraphs:
            for placeholder, value in replacement_data.items():
                self.replace_in_paragraph(paragraph, placeholder, value)

        self.replace_in_tables(document, replacement_data)
        self.replace_in_headers(document, replacement_data)
        self.replace_in_footers(document, replacement_data)
        self.replace_images(document, replacement_data)

        # Remove empty table rows
        self.remove_empty_table_rows(document)

    def replace_in_paragraph(self, paragraph, placeholder, value):
        """Mengganti placeholder di dalam paragraf."""
        if placeholder in paragraph.text and 'IMAGE' not in placeholder:
            paragraph.text = paragraph.text.replace(placeholder, value)
            for run in paragraph.runs:
                run.font.name = "Gordita Light"
                run.font.size = 9 * 12700

    def replace_in_tables(self, document, replacement_data):
        """Mengganti placeholder di dalam sel tabel."""
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in replacement_data.items():
                            self.replace_in_paragraph(paragraph, placeholder, value)

    def replace_in_headers(self, document, replacement_data):
        """Mengganti placeholder di dalam header."""
        for section in document.sections:
            header = section.header
            for paragraph in header.paragraphs:
                for placeholder, value in replacement_data.items():
                    self.replace_in_paragraph(paragraph, placeholder, value)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, value in replacement_data.items():
                                self.replace_in_paragraph(paragraph, placeholder, value)

    def replace_in_footers(self, document, replacement_data):
        """Mengganti placeholder di dalam footer."""
        for section in document.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                for placeholder, value in replacement_data.items():
                    self.replace_in_paragraph(paragraph, placeholder, value)
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, value in replacement_data.items():
                                self.replace_in_paragraph(paragraph, placeholder, value)

    def replace_images(self, document, replacement_data):
        """Mengganti placeholder gambar dengan gambar yang dipilih."""
        def process_paragraph(paragraph):
            placeholder = paragraph.text.strip()
            if placeholder in replacement_data:
                image_path = replacement_data[placeholder]
                if image_path and os.path.exists(image_path):
                    paragraph.clear()
                    # Check if it's a photography image (IMAGE3 to IMAGE12) for vertical orientation
                    if placeholder in ["[IMAGE3]", "[IMAGE4]", "[IMAGE5]", "[IMAGE6]", "[IMAGE7]", "[IMAGE8]", "[IMAGE9]", "[IMAGE10]", "[IMAGE11]", "[IMAGE12]"]:
                        # Smaller size for fitting 4 images on one page: height 3 inches, width 2 inches
                        paragraph.add_run().add_picture(image_path, width=Inches(2), height=Inches(3))
                    else:
                        # Other images: default size
                        paragraph.add_run().add_picture(image_path)
                else:
                    paragraph.clear()  # Remove placeholder if no valid image

        # Replace in main paragraphs
        for paragraph in document.paragraphs:
            process_paragraph(paragraph)

        # Replace in tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph(paragraph)

        # Replace in headers
        for section in document.sections:
            header = section.header
            for paragraph in header.paragraphs:
                process_paragraph(paragraph)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_paragraph(paragraph)

        # Replace in footers
        for section in document.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                process_paragraph(paragraph)
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_paragraph(paragraph)

    def remove_empty_table_rows(self, document):
        for table in document.tables:
            rows_to_remove = []
            for i, row in enumerate(table.rows):
                is_empty = True
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            is_empty = False
                            break
                    if not is_empty:
                        break
                if is_empty:
                    rows_to_remove.append(i)
            for i in reversed(rows_to_remove):
                table._tbl.remove(table.rows[i]._tr)

    def save_document(self, document):
        output_filename = f"Generated_Anexo_II_{date.today().strftime('%d_%m_%Y')}.docx"
        file_path, _ = QFileDialog.getSaveFileName(self.parent_app, "Guardar documento como...", output_filename, "Word Documents (*.docx);;All Files (*)", options=QFileDialog.Options())
        if not file_path:
            QMessageBox.information(self.parent_app, "Cancelado", "El almacenamiento fue cancelado por el usuario.")
            return
        if not file_path.lower().endswith('.docx'):
            file_path += '.docx'

        try:
            document.save(file_path)
            QMessageBox.information(
                self.parent_app,
                "¡Listo!",
                f"El documento de Word se creó y se guardó con éxito como:\n{file_path}"
            )
        except Exception as e:
            QMessageBox.critical(self.parent_app, "Error al guardar el archivo", f"Error al guardar el documento: {e}")
